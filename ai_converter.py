#!/usr/bin/env python3
"""
AI-Powered Iterative Template Converter
Uses AI (OpenAI or Anthropic) to automatically improve template conversions
"""

import os
import json
from typing import Dict, List, Tuple, Optional
from docx import Document
import difflib


class AIConverter:
    """
    Manages AI-powered iterative conversion improvements
    """

    def __init__(self, provider: str = 'openai', api_key: str = '', use_learning: bool = True):
        """
        Initialize AI converter

        Args:
            provider: 'openai' or 'anthropic'
            api_key: API key for the provider
            use_learning: Whether to use the learning cache
        """
        self.provider = provider
        self.api_key = api_key
        self.client = None
        self.learner = None

        if provider == 'openai':
            import openai
            self.client = openai.OpenAI(api_key=api_key)
        elif provider == 'anthropic':
            import anthropic
            self.client = anthropic.Anthropic(api_key=api_key)

        # Initialize learning system
        if use_learning:
            from conversion_learner import ConversionLearner
            self.learner = ConversionLearner()

    def extract_document_text(self, docx_path: str) -> str:
        """Extract all text from a Word document"""
        try:
            doc = Document(docx_path)
            text_parts = []

            for para in doc.paragraphs:
                text_parts.append(para.text)

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text_parts.append(cell.text)

            return '\n'.join(text_parts)
        except Exception as e:
            print(f"Error extracting text: {e}")
            return ""

    def extract_docx_templater_errors(self, docx_path: str) -> List[Dict]:
        """
        Extract DocX Templater error comments from a Word document

        Returns:
            list of dicts with:
                - error_text: The error message from the comment
                - location: Where the error appears (paragraph text near it)
        """
        import zipfile
        import xml.etree.ElementTree as ET

        errors = []

        try:
            with zipfile.ZipFile(docx_path, 'r') as zip_ref:
                # Read comments.xml if it exists
                try:
                    comments_xml = zip_ref.read('word/comments.xml').decode('utf-8')
                except KeyError:
                    return []  # No comments in document

                # Parse XML
                root = ET.fromstring(comments_xml)

                # Define namespace
                namespaces = {
                    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
                }

                # Find all comments
                for comment in root.findall('.//w:comment', namespaces):
                    # Extract comment text
                    comment_text_parts = []
                    for t in comment.findall('.//w:t', namespaces):
                        if t.text:
                            comment_text_parts.append(t.text)

                    comment_text = ''.join(comment_text_parts)

                    # Check if it looks like a DocX Templater error
                    if any(keyword in comment_text.lower() for keyword in
                           ['error', 'syntax', 'parse', 'tag', 'template', 'invalid']):
                        errors.append({
                            'error_text': comment_text,
                            'comment_id': comment.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id', 'unknown')
                        })

        except Exception as e:
            print(f"Error extracting comments: {e}")

        return errors

    def compare_documents(self, v1_doc_path: str, v2_doc_path: str) -> Dict:
        """
        Compare two documents and return differences

        Returns:
            dict with:
                - similarity_ratio: float 0-1
                - differences: list of diffs
                - missing_in_v2: list of content missing from v2
                - extra_in_v2: list of extra content in v2
        """
        v1_text = self.extract_document_text(v1_doc_path)
        v2_text = self.extract_document_text(v2_doc_path)

        # Split into lines for comparison
        v1_lines = v1_text.split('\n')
        v2_lines = v2_text.split('\n')

        # Calculate similarity
        similarity = difflib.SequenceMatcher(None, v1_text, v2_text).ratio()

        # Get differences
        diff = list(difflib.unified_diff(v1_lines, v2_lines, lineterm=''))

        # Analyze missing/extra content
        missing_in_v2 = []
        extra_in_v2 = []

        for line in diff:
            if line.startswith('- '):
                missing_in_v2.append(line[2:])
            elif line.startswith('+ '):
                extra_in_v2.append(line[2:])

        # Check for DocX Templater errors in V2
        v2_errors = self.extract_docx_templater_errors(v2_doc_path)

        return {
            'similarity_ratio': similarity,
            'differences': diff[:100],  # Limit to first 100 diff lines
            'missing_in_v2': missing_in_v2[:50],
            'extra_in_v2': extra_in_v2[:50],
            'v1_line_count': len(v1_lines),
            'v2_line_count': len(v2_lines),
            'v2_template_errors': v2_errors
        }

    def analyze_template_xml(self, template_path: str) -> str:
        """Extract relevant XML from template for AI analysis"""
        import zipfile
        try:
            with zipfile.ZipFile(template_path, 'r') as zip_ref:
                # Read document.xml
                xml_content = zip_ref.read('word/document.xml').decode('utf-8')
                # Truncate if too long (AI has token limits)
                if len(xml_content) > 50000:
                    xml_content = xml_content[:50000] + '\n... [truncated]'
                return xml_content
        except Exception as e:
            return f"Error reading XML: {e}"

    def suggest_mapping_improvements(
        self,
        v1_template_xml: str,
        v2_template_xml: str,
        document_comparison: Dict,
        current_mappings: List[Dict],
        iteration: int
    ) -> Dict:
        """
        Use AI to suggest mapping improvements based on document comparison

        Returns:
            dict with:
                - suggested_changes: list of mapping changes
                - reasoning: explanation from AI
                - confidence: float 0-1
        """
        # Build prompt for AI
        prompt = self._build_analysis_prompt(
            v1_template_xml,
            v2_template_xml,
            document_comparison,
            current_mappings,
            iteration
        )

        # Call AI
        if self.provider == 'openai':
            return self._call_openai(prompt)
        elif self.provider == 'anthropic':
            return self._call_anthropic(prompt)
        else:
            raise ValueError(f"Unknown provider: {self.provider}")

    def _build_analysis_prompt(
        self,
        v1_xml: str,
        v2_xml: str,
        comparison: Dict,
        mappings: List[Dict],
        iteration: int
    ) -> str:
        """Build the AI prompt for analysis"""

        # Truncate XMLs if needed
        v1_xml_short = v1_xml[:10000] + '\n... [truncated]' if len(v1_xml) > 10000 else v1_xml
        v2_xml_short = v2_xml[:10000] + '\n... [truncated]' if len(v2_xml) > 10000 else v2_xml

        prompt = f"""You are an expert at converting Microsoft Word Mail Merge templates to DocX Templater format.

ITERATION: {iteration}

CONTEXT:
We are converting a v1 (Mail Merge) template to v2 (DocX Templater) format.
The generated documents are {comparison['similarity_ratio']*100:.1f}% similar.

DOCUMENT COMPARISON:
- V1 document has {comparison['v1_line_count']} lines
- V2 document has {comparison['v2_line_count']} lines
- Missing from V2 ({len(comparison['missing_in_v2'])} items):
{chr(10).join(comparison['missing_in_v2'][:20])}

- Extra in V2 ({len(comparison['extra_in_v2'])} items):
{chr(10).join(comparison['extra_in_v2'][:20])}

CURRENT MAPPINGS (sample):
{json.dumps(mappings[:20], indent=2)}

V1 TEMPLATE XML (sample):
{v1_xml_short}

V2 TEMPLATE XML (sample):
{v2_xml_short}

YOUR TASK:
Analyze the differences and suggest SPECIFIC mapping changes to make V2 output match V1 output.

Focus on:
1. Unmapped fields (content missing from V2)
2. Incorrect loop structures
3. Missing conditionals
4. Field path errors

Return JSON with this structure:
{{
    "suggested_changes": [
        {{
            "v1_field": "=field_name",
            "current_v2": "{{current_path}}",
            "suggested_v2": "{{better_path}}",
            "reasoning": "why this change"
        }}
    ],
    "overall_assessment": "brief summary",
    "confidence": 0.8
}}

Return ONLY valid JSON, no other text.
"""
        return prompt

    def _call_openai(self, prompt: str) -> Dict:
        """Call OpenAI API"""
        try:
            response = self.client.chat.completions.create(
                model="gpt-4-turbo-preview",
                messages=[
                    {"role": "system", "content": "You are an expert template conversion assistant. Always respond with valid JSON."},
                    {"role": "user", "content": prompt}
                ],
                response_format={"type": "json_object"},
                temperature=0.3
            )

            result_text = response.choices[0].message.content
            return json.loads(result_text)

        except Exception as e:
            print(f"OpenAI API error: {e}")
            return {
                "suggested_changes": [],
                "overall_assessment": f"API Error: {str(e)}",
                "confidence": 0.0
            }

    def _call_anthropic(self, prompt: str) -> Dict:
        """Call Anthropic API"""
        try:
            response = self.client.messages.create(
                model="claude-3-5-sonnet-20241022",
                max_tokens=4000,
                temperature=0.3,
                messages=[
                    {"role": "user", "content": prompt}
                ]
            )

            result_text = response.content[0].text

            # Extract JSON from response (Claude might wrap it in markdown)
            if '```json' in result_text:
                start = result_text.find('```json') + 7
                end = result_text.find('```', start)
                result_text = result_text[start:end].strip()
            elif '```' in result_text:
                start = result_text.find('```') + 3
                end = result_text.find('```', start)
                result_text = result_text[start:end].strip()

            return json.loads(result_text)

        except Exception as e:
            print(f"Anthropic API error: {e}")
            return {
                "suggested_changes": [],
                "overall_assessment": f"API Error: {str(e)}",
                "confidence": 0.0
            }

    def fix_syntax_errors(
        self,
        template_xml: str,
        syntax_errors: List[Dict],
        mappings: List[Dict],
        merge_data_context: Optional[Dict] = None
    ) -> Dict:
        """
        Use AI to fix DocX Templater syntax errors in template XML

        Args:
            template_xml: The XML content with errors
            syntax_errors: List of detected errors
            mappings: Current field mappings
            merge_data_context: Optional dict with v1/v2 merge data structures for context

        Returns:
            dict with:
                - fixed_xml: Corrected XML content
                - fixes_applied: List of fixes made
                - reasoning: Explanation
        """
        if not syntax_errors:
            return {
                'fixed_xml': template_xml,
                'fixes_applied': [],
                'reasoning': 'No syntax errors to fix'
            }

        # Check if we have learned fixes for these errors
        known_fixes = []
        if self.learner:
            for error in syntax_errors:
                learned = self.learner.get_known_syntax_fixes(error['error_text'])
                if learned:
                    known_fixes.extend(learned)

        # Build prompt for syntax error fixing
        error_list = '\n'.join([f"- {err['error_text']}" for err in syntax_errors[:10]])

        known_fixes_text = ""
        if known_fixes:
            known_fixes_text = "\n\nKNOWN FIXES FROM PAST SUCCESSES:\n"
            for fix in known_fixes[:5]:
                known_fixes_text += f"- {fix['fix_pattern']} (confidence: {fix['confidence']*100:.0f}%)\n"
                known_fixes_text += f"  Pattern: {fix['search_regex']}\n"
                known_fixes_text += f"  Replacement: {fix['replacement']}\n"

        # Add merge data context if available
        merge_data_text = ""
        if merge_data_context:
            v1_structure = merge_data_context.get('v1_structure_sample', '')
            v2_structure = merge_data_context.get('v2_structure_sample', '')

            if v1_structure or v2_structure:
                merge_data_text = "\n\nMERGE DATA STRUCTURE CONTEXT:\n"
                merge_data_text += "IMPORTANT: V1 and V2 merge data have DIFFERENT structures.\n"
                merge_data_text += "Fields must use V2 structure paths, not V1 paths.\n\n"

                if v1_structure:
                    merge_data_text += "V1 Merge Data Structure (OLD - for reference only):\n"
                    merge_data_text += f"{v1_structure}\n\n"

                if v2_structure:
                    merge_data_text += "V2 Merge Data Structure (NEW - use these paths):\n"
                    merge_data_text += f"{v2_structure}\n\n"

                merge_data_text += "When fixing field paths, ensure they match the V2 structure above.\n"

        prompt = f"""You are an expert at fixing DocX Templater syntax errors in Word document XML.

REFERENCE DOCUMENTATION (DocX Templater):
- Tags: Use {{field}} for simple fields, {{#loop}}...{{/loop}} for loops
- Conditions: {{#condition}}...{{/condition}} or {{^condition}}...{{/condition}} for inverted
- Loops: {{#items}}{{name}}{{/items}} - opening and closing tags MUST match exactly
- Nested paths: Use dot notation like {{project.name}} or {{customer.address.city}}
- Raw XML: Use {{@rawXML}} to insert XML content
- Documentation: https://docxtemplater.com/docs/tag-types/

SYNTAX ERRORS DETECTED:
{error_list}

CURRENT MAPPINGS:
{json.dumps(mappings[:30], indent=2)}
{merge_data_text}{known_fixes_text}
TEMPLATE XML (sample):
{template_xml[:15000]}

YOUR TASK:
Fix the DocX Templater syntax errors in the template XML. Common issues:
1. Unclosed loops: Every {{#loop}} needs a {{/loop}} with EXACT same name
2. Mismatched tags: Opening tag "{{#project.items}}" must close with "{{/project.items}}" (not "{{/items}}")
3. Unopened loops: {{/loop}} without a matching {{#loop}}
4. Invalid field paths: Must match the data structure (use dot notation for nested fields)
5. Text runs: DocX Templater tags in Word XML are often split across <w:t> elements - search carefully

Return JSON with this structure:
{{
    "fixes": [
        {{
            "error": "description of error",
            "location": "approximate location in XML",
            "fix": "what to change",
            "search_pattern": "text to find (regex)",
            "replacement": "text to replace with"
        }}
    ],
    "reasoning": "overall explanation of fixes",
    "confidence": 0.8
}}

Return ONLY valid JSON, no other text.
"""

        # Call AI
        if self.provider == 'openai':
            response = self._call_openai(prompt)
        elif self.provider == 'anthropic':
            response = self._call_anthropic(prompt)
        else:
            raise ValueError(f"Unknown provider: {self.provider}")

        # Apply fixes to XML
        fixed_xml = template_xml
        fixes_applied = []

        for fix in response.get('fixes', []):
            search_pattern = fix.get('search_pattern', '')
            replacement = fix.get('replacement', '')

            if search_pattern and replacement:
                import re
                try:
                    # Apply the fix
                    new_xml = re.sub(search_pattern, replacement, fixed_xml, flags=re.DOTALL)
                    if new_xml != fixed_xml:
                        fixed_xml = new_xml
                        fix_info = {
                            'error': fix.get('error'),
                            'fix': fix.get('fix')
                        }
                        fixes_applied.append(fix_info)

                        # Record successful fix in learning cache
                        if self.learner:
                            self.learner.record_syntax_fix(
                                error_pattern=fix.get('error', ''),
                                fix_pattern=fix.get('fix', ''),
                                search_regex=search_pattern,
                                replacement=replacement
                            )
                except Exception as e:
                    print(f"Failed to apply fix: {e}")

        return {
            'fixed_xml': fixed_xml,
            'fixes_applied': fixes_applied,
            'reasoning': response.get('reasoning', ''),
            'confidence': response.get('confidence', 0.0)
        }

    def iterative_convert(
        self,
        v1_template_path: str,
        v2_template_path: str,
        v1_document_path: str,
        v2_document_path: str,
        initial_mappings: List[Dict],
        max_iterations: int = 4,
        target_similarity: float = 0.95,
        callback=None
    ) -> Dict:
        """
        Iteratively improve conversion using AI

        Args:
            v1_template_path: Path to v1 template
            v2_template_path: Path to v2 template (will be regenerated)
            v1_document_path: Path to generated v1 document
            v2_document_path: Path to generated v2 document
            initial_mappings: Starting mappings
            max_iterations: Max number of improvement iterations
            target_similarity: Stop if we reach this similarity (0-1)
            callback: Optional function(iteration, result) for progress updates

        Returns:
            dict with iteration history and final mappings
        """
        iteration_history = []
        current_mappings = initial_mappings.copy()

        # Get template XMLs once
        v1_xml = self.analyze_template_xml(v1_template_path)
        v2_xml = self.analyze_template_xml(v2_template_path)

        for iteration in range(1, max_iterations + 1):
            print(f"\n{'='*60}")
            print(f"Iteration {iteration}/{max_iterations}")
            print(f"{'='*60}")

            # Compare documents
            comparison = self.compare_documents(v1_document_path, v2_document_path)

            print(f"📊 Document Comparison:")
            print(f"   Similarity: {comparison['similarity_ratio']*100:.1f}%")
            print(f"   V1 lines: {comparison['v1_line_count']}")
            print(f"   V2 lines: {comparison['v2_line_count']}")
            print(f"   Missing from V2: {len(comparison['missing_in_v2'])} items")
            print(f"   Extra in V2: {len(comparison['extra_in_v2'])} items")

            # Report DocX Templater errors if found
            if comparison.get('v2_template_errors'):
                print(f"\n⚠️  DocX Templater Syntax Errors Found: {len(comparison['v2_template_errors'])}")
                for i, error in enumerate(comparison['v2_template_errors'][:5], 1):
                    print(f"   Error {i}: {error['error_text'][:100]}")

            if callback:
                callback(iteration, {
                    'type': 'comparison',
                    'similarity': comparison['similarity_ratio'],
                    'errors': comparison.get('v2_template_errors', [])
                })

            # Check if we've reached target
            if comparison['similarity_ratio'] >= target_similarity:
                print(f"✓ Target similarity reached: {comparison['similarity_ratio']*100:.1f}%")
                iteration_history.append({
                    'iteration': iteration,
                    'similarity': comparison['similarity_ratio'],
                    'changes_applied': 0,
                    'status': 'target_reached'
                })
                break

            # Ask AI for suggestions
            print(f"Asking AI for mapping improvements...")
            ai_response = self.suggest_mapping_improvements(
                v1_xml,
                v2_xml,
                comparison,
                current_mappings,
                iteration
            )

            if callback:
                callback(iteration, {
                    'type': 'ai_response',
                    'response': ai_response
                })

            # Apply suggested changes
            changes_applied = 0
            for change in ai_response.get('suggested_changes', []):
                # Update mapping
                v1_field = change['v1_field']
                suggested_v2 = change['suggested_v2']

                # Find and update the mapping
                for mapping in current_mappings:
                    if mapping.get('v1_field') == v1_field:
                        mapping['v2_field'] = suggested_v2
                        changes_applied += 1
                        break

            iteration_history.append({
                'iteration': iteration,
                'similarity': comparison['similarity_ratio'],
                'changes_suggested': len(ai_response.get('suggested_changes', [])),
                'changes_applied': changes_applied,
                'ai_assessment': ai_response.get('overall_assessment', ''),
                'ai_confidence': ai_response.get('confidence', 0.0),
                'status': 'continued'
            })

            if changes_applied == 0:
                print("No more changes to apply. Stopping.")
                iteration_history[-1]['status'] = 'no_changes'
                break

            # TODO: Here we would need to reconvert the template with new mappings
            # and regenerate the v2 document. This requires integration with
            # the conversion logic.

            print(f"Applied {changes_applied} mapping changes")

        return {
            'iterations': iteration_history,
            'final_mappings': current_mappings,
            'final_similarity': iteration_history[-1]['similarity'] if iteration_history else 0.0
        }
